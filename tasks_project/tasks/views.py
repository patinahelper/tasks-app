from django.shortcuts import render, get_object_or_404, redirect
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.http import JsonResponse
from django.views.decorators.http import require_POST
from django.utils import timezone
import json
from .models import Project, Task, ChatMessage, Incident, TaskUpdate, SubTask
from .forms import TaskForm, ProjectForm, ChatMessageForm, IncidentForm, TaskUpdateForm, SubTaskForm

def chat_view(request):
    """Chat interface with BMO"""
    if request.method == 'POST':
        form = ChatMessageForm(request.POST)
        if form.is_valid():
            message = form.save(commit=False)
            message.sender = 'user'
            message.save()
            return redirect('chat')
    else:
        form = ChatMessageForm()
    
    # Get recent messages (last 50) - ordered by timestamp
    messages = ChatMessage.objects.order_by('-timestamp')[:50]
    # Reverse so oldest first for display
    messages = list(reversed(messages))
    
    # Get last message ID for polling
    last_message_id = messages[-1].id if messages else 0
    
    context = {
        'messages': messages,
        'form': form,
        'unread_count': ChatMessage.objects.filter(sender='agent', is_read=False).count(),
        'last_message_id': last_message_id,
    }
    return render(request, 'tasks/chat.html', context)

@require_POST
def chat_api(request):
    """AJAX endpoint to send message"""
    data = json.loads(request.body)
    content = data.get('message', '').strip()
    
    if content:
        ChatMessage.objects.create(
            content=content,
            sender='user',
            timestamp=timezone.now()
        )
        return JsonResponse({'success': True})
    
    return JsonResponse({'success': False, 'error': 'Empty message'})

def chat_poll(request):
    """AJAX endpoint to poll for new messages"""
    last_id = request.GET.get('last_id')
    
    if last_id:
        messages = ChatMessage.objects.filter(id__gt=last_id).values('id', 'content', 'sender', 'timestamp')
    else:
        messages = ChatMessage.objects.all().values('id', 'content', 'sender', 'timestamp')
    
    # Mark agent messages as read
    ChatMessage.objects.filter(sender='agent', is_read=False).update(is_read=True)
    
    return JsonResponse({
        'messages': list(messages),
        'unread_count': ChatMessage.objects.filter(sender='agent', is_read=False).count(),
    })

def dashboard(request):
    """Main dashboard view"""
    context = {
        'total_tasks': Task.objects.count(),
        'tasks_by_status': {
            'backlog': Task.objects.filter(status='backlog').count(),
            'todo': Task.objects.filter(status='todo').count(),
            'in_progress': Task.objects.filter(status='in_progress').count(),
            'review': Task.objects.filter(status='review').count(),
            'done': Task.objects.filter(status='done').count(),
        },
        'urgent_tasks': Task.objects.filter(priority=4, status__in=['backlog', 'todo', 'in_progress']).order_by('due_date')[:5],
        'recent_tasks': Task.objects.order_by('-updated_at')[:10],
        'projects': Project.objects.filter(is_active=True),
    }
    return render(request, 'tasks/dashboard.html', context)

def kanban_board(request):
    """Kanban board view"""
    project_id = request.GET.get('project')
    tasks = Task.objects.all()
    
    if project_id:
        tasks = tasks.filter(project_id=project_id)
    
    context = {
        'columns': [
            {'id': 'backlog', 'name': 'Backlog', 'tasks': tasks.filter(status='backlog')},
            {'id': 'todo', 'name': 'To Do', 'tasks': tasks.filter(status='todo')},
            {'id': 'in_progress', 'name': 'In Progress', 'tasks': tasks.filter(status='in_progress')},
            {'id': 'review', 'name': 'Review', 'tasks': tasks.filter(status='review')},
            {'id': 'done', 'name': 'Done', 'tasks': tasks.filter(status='done')},
        ],
        'projects': Project.objects.filter(is_active=True),
        'current_project': int(project_id) if project_id else None,
    }
    return render(request, 'tasks/kanban.html', context)

@require_POST
def task_move(request):
    """AJAX endpoint to move task between columns"""
    data = json.loads(request.body)
    task_id = data.get('task_id')
    new_status = data.get('status')
    
    task = get_object_or_404(Task, pk=task_id)
    task.status = new_status
    if new_status == 'done':
        from django.utils import timezone
        task.completed_at = timezone.now()
    task.save()
    
    return JsonResponse({'success': True})

class ProjectListView(ListView):
    model = Project
    template_name = 'tasks/project_list.html'
    context_object_name = 'projects'

class ProjectDetailView(DetailView):
    model = Project
    template_name = 'tasks/project_detail.html'
    
    def get_context_data(self, **kwargs):
        from django.utils import timezone
        context = super().get_context_data(**kwargs)
        context['tasks'] = self.object.tasks.order_by('-priority', 'due_date')
        context['today'] = timezone.now().date()
        # Task counts for stats
        context['total_tasks'] = self.object.tasks.count()
        context['in_progress_count'] = self.object.tasks.filter(status='in_progress').count()
        context['completed_count'] = self.object.tasks.filter(status='done').count()
        return context

class ProjectCreateView(CreateView):
    model = Project
    form_class = ProjectForm
    template_name = 'tasks/project_form.html'
    success_url = reverse_lazy('project_list')

class ProjectUpdateView(UpdateView):
    model = Project
    form_class = ProjectForm
    template_name = 'tasks/project_form.html'
    success_url = reverse_lazy('project_list')

class TaskListView(ListView):
    model = Task
    template_name = 'tasks/task_list.html'
    context_object_name = 'tasks'
    paginate_by = 20

    def get_queryset(self):
        from django.utils import timezone
        from datetime import timedelta
        queryset = super().get_queryset()
        status = self.request.GET.get('status')
        project = self.request.GET.get('project')
        priority = self.request.GET.get('priority')
        due = self.request.GET.get('due')

        if status:
            queryset = queryset.filter(status=status)
        if project:
            queryset = queryset.filter(project_id=project)
        if priority:
            queryset = queryset.filter(priority=priority)

        # Due date filters
        today = timezone.now().date()
        if due == 'today':
            queryset = queryset.filter(due_date=today)
        elif due == 'week':
            week_end = today + timedelta(days=7)
            queryset = queryset.filter(due_date__gte=today, due_date__lte=week_end)
        elif due == 'overdue':
            queryset = queryset.filter(due_date__lt=today, status__in=['backlog', 'todo', 'in_progress'])

        return queryset

    def get_context_data(self, **kwargs):
        from django.utils import timezone
        from datetime import timedelta
        context = super().get_context_data(**kwargs)
        context['projects'] = Project.objects.filter(is_active=True)
        context['status_choices'] = Task.STATUS_CHOICES
        context['priority_choices'] = Task.PRIORITY_CHOICES
        context['today'] = timezone.now().date()
        context['week_end'] = timezone.now().date() + timedelta(days=7)

        # Due date counts for filter badges
        today = timezone.now().date()
        week_end = today + timedelta(days=7)
        base_qs = Task.objects.exclude(status='done')
        context['due_today_count'] = base_qs.filter(due_date=today).count()
        context['due_this_week_count'] = base_qs.filter(due_date__gte=today, due_date__lte=week_end).count()
        context['overdue_count'] = base_qs.filter(due_date__lt=today).count()

        # Current filter
        context['current_due'] = self.request.GET.get('due', '')
        return context

class TaskDetailView(DetailView):
    model = Task
    template_name = 'tasks/task_detail.html'
    
    def get_context_data(self, **kwargs):
        from django.utils import timezone
        context = super().get_context_data(**kwargs)
        context['update_form'] = TaskUpdateForm()
        context['updates'] = self.object.updates.all()
        context['subtasks'] = self.object.subtasks.all()
        context['subtask_form'] = SubTaskForm()
        context['today'] = timezone.now().date()
        return context

class TaskCreateView(CreateView):
    model = Task
    form_class = TaskForm
    template_name = 'tasks/task_form.html'
    
    def get_initial(self):
        initial = super().get_initial()
        # Pre-select project if passed in URL
        project_id = self.request.GET.get('project')
        if project_id:
            initial['project'] = project_id
        return initial
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Remember where we came from for redirect
        context['return_to_project'] = self.request.GET.get('project')
        return context
    
    def get_success_url(self):
        # Redirect back to project if we came from there
        return_to_project = self.request.GET.get('project')
        if return_to_project:
            return reverse_lazy('project_detail', kwargs={'pk': return_to_project})
        return reverse_lazy('kanban')

class TaskUpdateView(UpdateView):
    model = Task
    form_class = TaskForm
    template_name = 'tasks/task_form.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Remember where we came from for redirect
        context['return_to_project'] = self.request.GET.get('return_to_project')
        return context
    
    def get_success_url(self):
        # Check if we should return to project
        return_to_project = self.request.GET.get('return_to_project')
        if return_to_project:
            return reverse_lazy('project_detail', kwargs={'pk': return_to_project})
        return reverse_lazy('task_detail', kwargs={'pk': self.object.pk})

class TaskDeleteView(DeleteView):
    model = Task
    success_url = reverse_lazy('kanban')
    template_name = 'tasks/task_confirm_delete.html'


class TaskUpdateCreateView(CreateView):
    model = TaskUpdate
    form_class = TaskUpdateForm
    template_name = 'tasks/taskupdate_form.html'
    
    def dispatch(self, request, *args, **kwargs):
        self.task = get_object_or_404(Task, pk=self.kwargs['task_pk'])
        return super().dispatch(request, *args, **kwargs)
    
    def form_valid(self, form):
        form.instance.task = self.task
        return super().form_valid(form)
    
    def get_success_url(self):
        return reverse_lazy('task_detail', kwargs={'pk': self.task.pk})
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['task'] = self.task
        return context


class SubTaskCreateView(CreateView):
    model = SubTask
    form_class = SubTaskForm
    template_name = 'tasks/subtask_form.html'
    
    def dispatch(self, request, *args, **kwargs):
        self.parent_task = get_object_or_404(Task, pk=self.kwargs['task_pk'])
        return super().dispatch(request, *args, **kwargs)
    
    def get_initial(self):
        initial = super().get_initial()
        # Pre-select parent task
        initial['parent_task'] = self.parent_task
        return initial
    
    def form_valid(self, form):
        form.instance.parent_task = self.parent_task
        return super().form_valid(form)
    
    def get_success_url(self):
        return reverse_lazy('task_detail', kwargs={'pk': self.parent_task.pk})
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['parent_task'] = self.parent_task
        return context


class SubTaskUpdateView(UpdateView):
    model = SubTask
    form_class = SubTaskForm
    template_name = 'tasks/subtask_form.html'
    
    def get_success_url(self):
        return reverse_lazy('task_detail', kwargs={'pk': self.object.parent_task.pk})


class SubTaskDeleteView(DeleteView):
    model = SubTask
    template_name = 'tasks/subtask_confirm_delete.html'
    
    def get_success_url(self):
        return reverse_lazy('task_detail', kwargs={'pk': self.object.parent_task.pk})


# ========== Weekly Report Views ==========

def weekly_report(request):
    """Weekly report view for Mohsen"""
    from django.utils import timezone
    from datetime import timedelta
    
    # Get date range
    period = request.GET.get('period', '7')
    debug = request.GET.get('debug') == '1'
    
    if period == 'last_week':
        # Last week: Monday to Sunday of previous week
        today = timezone.now().date()
        # Find last Monday (today.weekday() gives 0=Monday, so we go back 7 + weekday days)
        days_since_monday = today.weekday()
        last_sunday = today - timedelta(days=days_since_monday + 1)
        last_monday = last_sunday - timedelta(days=6)
        start_date = last_monday
        end_date = last_sunday
        days = 'last_week'
    else:
        days = int(period)
        end_date = timezone.now().date()
        start_date = end_date - timedelta(days=days)
    
    # Tasks completed this week
    # Use timezone-aware datetime comparison instead of just date
    start_datetime = timezone.make_aware(timezone.datetime.combine(start_date, timezone.datetime.min.time()))
    end_datetime = timezone.make_aware(timezone.datetime.combine(end_date, timezone.datetime.max.time()))
    
    completed_tasks = Task.objects.filter(
        completed_at__gte=start_datetime,
        completed_at__lte=end_datetime
    ).select_related('project').order_by('project__category', '-completed_at')
    
    # Debug info
    debug_info = None
    if debug:
        all_done = Task.objects.filter(status='done')
        debug_info = {
            'start_date': start_date,
            'end_date': end_date,
            'start_datetime': start_datetime,
            'end_datetime': end_datetime,
            'completed_count': completed_tasks.count(),
            'all_done_count': all_done.count(),
            'all_done_tasks': [(t.title, t.completed_at) for t in all_done[:5]],
        }
    
    # Tasks in progress
    in_progress_tasks = Task.objects.filter(
        status='in_progress'
    ).select_related('project').order_by('project__category', '-priority')
    
    # Upcoming deadlines (next 7 days)
    upcoming_tasks = Task.objects.filter(
        due_date__gte=end_date,
        due_date__lte=end_date + timedelta(days=7),
        status__in=['backlog', 'todo', 'in_progress']
    ).select_related('project').order_by('due_date')
    
    # Incidents
    new_incidents = Incident.objects.filter(
        date_reported__gte=start_date,
        date_reported__lte=end_date
    ).order_by('-severity', '-date_reported')
    
    open_incidents = Incident.objects.filter(
        status__in=['open', 'in_progress']
    ).order_by('-severity', '-date_reported')
    
    # Group by category
    categories = ['safety', 'operational', 'workshop', 'general']
    
    def group_by_category(tasks):
        grouped = {cat: [] for cat in categories}
        for task in tasks:
            cat = task.project.category if task.project else 'general'
            grouped[cat].append(task)
        return grouped
    
    context = {
        'start_date': start_date,
        'end_date': end_date,
        'days': days,
        'period': period,
        'completed_by_category': group_by_category(completed_tasks),
        'in_progress_by_category': group_by_category(in_progress_tasks),
        'upcoming_tasks': upcoming_tasks,
        'new_incidents': new_incidents,
        'open_incidents': open_incidents,
        'category_labels': dict(Project.CATEGORY_CHOICES),
        'debug_info': debug_info,
    }
    return render(request, 'tasks/weekly_report.html', context)


def weekly_report_email(request):
    """Plain text version for easy email copy/paste"""
    from django.utils import timezone
    from datetime import timedelta
    
    period = request.GET.get('period', '7')
    
    if period == 'last_week':
        today = timezone.now().date()
        days_since_monday = today.weekday()
        last_sunday = today - timedelta(days=days_since_monday + 1)
        last_monday = last_sunday - timedelta(days=6)
        start_date = last_monday
        end_date = last_sunday
    else:
        days = int(period)
        end_date = timezone.now().date()
        start_date = end_date - timedelta(days=days)
    
    # Same queries as weekly_report - use timezone-aware datetime
    start_datetime = timezone.make_aware(timezone.datetime.combine(start_date, timezone.datetime.min.time()))
    end_datetime = timezone.make_aware(timezone.datetime.combine(end_date, timezone.datetime.max.time()))
    
    completed_tasks = Task.objects.filter(
        completed_at__gte=start_datetime,
        completed_at__lte=end_datetime
    ).select_related('project').order_by('project__category', '-completed_at')
    
    in_progress_tasks = Task.objects.filter(
        status='in_progress'
    ).select_related('project').order_by('project__category', '-priority')
    
    upcoming_tasks = Task.objects.filter(
        due_date__gte=end_date,
        due_date__lte=end_date + timedelta(days=7),
        status__in=['backlog', 'todo', 'in_progress']
    ).select_related('project').order_by('due_date')
    
    new_incidents = Incident.objects.filter(
        date_reported__gte=start_date,
        date_reported__lte=end_date
    ).order_by('-severity', '-date_reported')
    
    open_incidents = Incident.objects.filter(
        status__in=['open', 'in_progress']
    ).order_by('-severity', '-date_reported')
    
    # Build plain text report
    lines = [
        f"WEEKLY REPORT - {start_date.strftime('%d %b')} to {end_date.strftime('%d %b %Y')}",
        "=" * 50,
        "",
    ]
    
    category_labels = dict(Project.CATEGORY_CHOICES)
    categories = ['safety', 'operational', 'workshop', 'general']
    
    # Completed tasks
    lines.append("COMPLETED THIS WEEK")
    lines.append("-" * 30)
    for cat in categories:
        cat_tasks = [t for t in completed_tasks if (t.project.category if t.project else 'general') == cat]
        if cat_tasks:
            lines.append(f"\n{category_labels[cat].upper()}")
            for task in cat_tasks:
                proj = f" [{task.project.name}]" if task.project else ""
                lines.append(f"  ✓ {task.title}{proj}")
    if not completed_tasks:
        lines.append("  No tasks completed this week.")
    lines.append("")
    
    # In progress
    lines.append("IN PROGRESS")
    lines.append("-" * 30)
    for cat in categories:
        cat_tasks = [t for t in in_progress_tasks if (t.project.category if t.project else 'general') == cat]
        if cat_tasks:
            lines.append(f"\n{category_labels[cat].upper()}")
            for task in cat_tasks:
                proj = f" [{task.project.name}]" if task.project else ""
                lines.append(f"  → {task.title}{proj}")
    if not in_progress_tasks:
        lines.append("  No tasks in progress.")
    lines.append("")
    
    # Upcoming deadlines
    if upcoming_tasks:
        lines.append("UPCOMING DEADLINES (Next 7 Days)")
        lines.append("-" * 30)
        for task in upcoming_tasks:
            due = task.due_date.strftime('%d %b') if task.due_date else 'No date'
            proj = f" [{task.project.name}]" if task.project else ""
            lines.append(f"  • {task.title} (due {due}){proj}")
        lines.append("")
    
    # Incidents
    if new_incidents or open_incidents:
        lines.append("HAZARDS & INCIDENTS")
        lines.append("-" * 30)
        
        if new_incidents:
            lines.append("\nNew this week:")
            for inc in new_incidents:
                severity = "HIGH" if inc.severity == 3 else "MED" if inc.severity == 2 else "LOW"
                lines.append(f"  ⚠ [{severity}] {inc.title}")
                if inc.action_taken:
                    lines.append(f"     Action: {inc.action_taken}")
        
        open_not_new = [i for i in open_incidents if i not in new_incidents]
        if open_not_new:
            lines.append("\nStill open:")
            for inc in open_not_new:
                severity = "HIGH" if inc.severity == 3 else "MED" if inc.severity == 2 else "LOW"
                lines.append(f"  ⚠ [{severity}] {inc.title}")
                if inc.action_taken:
                    lines.append(f"     Action: {inc.action_taken}")
        lines.append("")
    
    report_text = "\n".join(lines)
    
    return render(request, 'tasks/weekly_report_email.html', {
        'report_text': report_text,
        'start_date': start_date,
        'end_date': end_date,
    })


def weekly_report_word(request):
    """Generate Word document for weekly report"""
    from django.utils import timezone
    from datetime import timedelta
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from django.http import HttpResponse
    import io
    
    period = request.GET.get('period', '7')
    
    if period == 'last_week':
        today = timezone.now().date()
        days_since_monday = today.weekday()
        last_sunday = today - timedelta(days=days_since_monday + 1)
        last_monday = last_sunday - timedelta(days=6)
        start_date = last_monday
        end_date = last_sunday
    else:
        days = int(period)
        end_date = timezone.now().date()
        start_date = end_date - timedelta(days=days)
    
    # Get data - use timezone-aware datetime
    start_datetime = timezone.make_aware(timezone.datetime.combine(start_date, timezone.datetime.min.time()))
    end_datetime = timezone.make_aware(timezone.datetime.combine(end_date, timezone.datetime.max.time()))
    
    completed_tasks = Task.objects.filter(
        completed_at__gte=start_datetime,
        completed_at__lte=end_datetime
    ).select_related('project').order_by('project__category', '-completed_at')
    
    in_progress_tasks = Task.objects.filter(
        status='in_progress'
    ).select_related('project').order_by('project__category', '-priority')
    
    upcoming_tasks = Task.objects.filter(
        due_date__gte=end_date,
        due_date__lte=end_date + timedelta(days=7),
        status__in=['backlog', 'todo', 'in_progress']
    ).select_related('project').order_by('due_date')
    
    new_incidents = Incident.objects.filter(
        date_reported__gte=start_date,
        date_reported__lte=end_date
    ).order_by('-severity', '-date_reported')
    
    open_incidents = Incident.objects.filter(
        status__in=['open', 'in_progress']
    ).order_by('-severity', '-date_reported')
    
    # Create document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Weekly Activity Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with name and date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Mitch Alexander - Laboratory Facilitator\n')
    run.font.size = Pt(12)
    run.font.bold = True
    run = subtitle.add_run(f'{start_date.strftime("%d %B %Y")} - {end_date.strftime("%d %B %Y")}')
    run.font.size = Pt(11)
    run.font.italic = True
    
    doc.add_paragraph()
    
    # Helper function to add category section
    category_labels = dict(Project.CATEGORY_CHOICES)
    categories = ['safety', 'operational', 'workshop', 'general']
    
    def add_category_section(doc, heading, tasks, show_project=True):
        doc.add_heading(heading, level=1)
        
        has_content = False
        for cat in categories:
            cat_tasks = [t for t in tasks if (t.project.category if t.project else 'general') == cat]
            if cat_tasks:
                has_content = True
                # Category subheading
                p = doc.add_paragraph()
                run = p.add_run(category_labels[cat])
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)  # Navy blue
                
                # Task list
                for task in cat_tasks:
                    p = doc.add_paragraph(style='List Bullet')
                    task_text = task.title
                    if show_project and task.project:
                        task_text += f' [{task.project.name}]'
                    p.add_run(task_text)
                    
                    # Add priority note for in-progress tasks
                    if hasattr(task, 'priority_label') and heading == 'Tasks In Progress':
                        p.add_run(f' - Priority: {task.priority_label}').font.italic = True
        
        if not has_content:
            doc.add_paragraph('No items to report.', style='List Bullet')
        
        doc.add_paragraph()
    
    # Completed tasks
    add_category_section(doc, 'Completed This Week', completed_tasks)
    
    # In Progress tasks
    add_category_section(doc, 'Tasks In Progress', in_progress_tasks)
    
    # Upcoming deadlines
    if upcoming_tasks:
        doc.add_heading('Upcoming Deadlines (Next 7 Days)', level=1)
        for task in upcoming_tasks:
            p = doc.add_paragraph(style='List Bullet')
            due_text = f'{task.title} (Due: {task.due_date.strftime("%d %b")})'
            if task.project:
                due_text += f' [{task.project.name}]'
            p.add_run(due_text)
        doc.add_paragraph()
    
    # Incidents
    if new_incidents or open_incidents:
        doc.add_heading('Hazards and Incidents', level=1)
        
        if new_incidents:
            p = doc.add_paragraph()
            run = p.add_run('New This Week')
            run.font.bold = True
            run.font.size = Pt(11)
            
            for inc in new_incidents:
                p = doc.add_paragraph(style='List Bullet')
                severity_text = 'HIGH' if inc.severity == 3 else 'MED' if inc.severity == 2 else 'LOW'
                run = p.add_run(f'[{severity_text}] ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) if inc.severity == 3 else RGBColor(0xFF, 0x80, 0x00) if inc.severity == 2 else RGBColor(0x00, 0x80, 0x00)
                p.add_run(f'{inc.title}')
                if inc.action_taken:
                    action_p = doc.add_paragraph(style='List Bullet 2')
                    action_p.add_run(f'Action taken: {inc.action_taken}').font.italic = True
        
        open_not_new = [i for i in open_incidents if i not in new_incidents]
        if open_not_new:
            p = doc.add_paragraph()
            run = p.add_run('Still Open')
            run.font.bold = True
            run.font.size = Pt(11)
            
            for inc in open_not_new:
                p = doc.add_paragraph(style='List Bullet')
                severity_text = 'HIGH' if inc.severity == 3 else 'MED' if inc.severity == 2 else 'LOW'
                run = p.add_run(f'[{severity_text}] ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) if inc.severity == 3 else RGBColor(0xFF, 0x80, 0x00) if inc.severity == 2 else RGBColor(0x00, 0x80, 0x00)
                p.add_run(f'{inc.title}')
                if inc.action_taken:
                    action_p = doc.add_paragraph(style='List Bullet 2')
                    action_p.add_run(f'Action taken: {inc.action_taken}').font.italic = True
        
        doc.add_paragraph()
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'Generated {timezone.localtime(timezone.now()).strftime("%d %b %Y at %H:%M")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    # Save to response
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="Weekly_Report_{end_date.strftime("%Y%m%d")}.docx"'
    
    return response


# ========== Incident Views ==========

class IncidentListView(ListView):
    model = Incident
    template_name = 'tasks/incident_list.html'
    context_object_name = 'incidents'
    paginate_by = 20

    def get_queryset(self):
        queryset = super().get_queryset()
        status = self.request.GET.get('status')
        severity = self.request.GET.get('severity')
        
        if status:
            queryset = queryset.filter(status=status)
        if severity:
            queryset = queryset.filter(severity=severity)
        
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['status_choices'] = Incident.STATUS_CHOICES
        context['severity_choices'] = Incident.SEVERITY_CHOICES
        context['open_count'] = Incident.objects.filter(status__in=['open', 'in_progress']).count()
        return context


class IncidentDetailView(DetailView):
    model = Incident
    template_name = 'tasks/incident_detail.html'


class IncidentCreateView(CreateView):
    model = Incident
    form_class = IncidentForm
    template_name = 'tasks/incident_form.html'
    success_url = reverse_lazy('incident_list')


class IncidentUpdateView(UpdateView):
    model = Incident
    form_class = IncidentForm
    template_name = 'tasks/incident_form.html'
    success_url = reverse_lazy('incident_list')


class IncidentDeleteView(DeleteView):
    model = Incident
    success_url = reverse_lazy('incident_list')
    template_name = 'tasks/incident_confirm_delete.html'
